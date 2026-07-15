"""Shared ingest adapters — turn ANY source into a raw table (list of rows) or plain
text, so the diversity of inputs lives here and the cleanup pipeline (dataclean.py) stays one
shared path. Used by `data-tidy` (clean tabular data) and `data-extract` (get data
out of documents). Lives at the toolkit-root `scripts/`; skills import it by adding
`../../scripts` to sys.path (run from the skill directory).

    from ingest import read_any, read_paste, read_text, list_sheets, list_pdf_tables, ocr_available
    rows, note = read_any("export.xlsx")       # or .csv/.tsv, a .pdf, a .docx, a .msg
    rows, note = read_any("book.xlsx", sheet="Q2")  # pick a tab in a multi-sheet workbook
    sheets     = list_sheets("book.xlsx")       # discover tabs (name/size/hidden/empty)
    rows, note = read_paste(pasted_text)        # markdown / tab / csv-ish pasted table
    text, note = read_text("form.pdf")          # document -> plain text (OCR fallback)
    cands      = list_pdf_tables("report.pdf")  # enumerate tables per page for selection

Multi-sheet workbooks: read_any/read_xlsx auto-select the single non-empty sheet; if several
hold data they raise `SheetSelectionRequired` (never silently read the 'active' tab) — the
caller lists `list_sheets()` and re-calls with `sheet=`.

Heavy deps (PyMuPDF, pdfplumber, python-docx, extract_msg, Tesseract) are imported LAZILY and
degrade with a clear message rather than crashing — most jobs (xlsx / csv / paste / digital
PDF) need none of them.

PDF tables: two engines, best result per page. pdfplumber (optional, preferred for messy /
borderless tables) and PyMuPDF (fast, ruled tables, + the OCR backbone) are both tried when
installed; each page keeps the higher-scoring extraction. With pdfplumber absent it's PyMuPDF
only — same as before.

OCR (scanned PDFs): tried ONLY when a PDF page has no text layer, and ONLY via LOCAL
Tesseract — never a cloud OCR, so no third-party service ever sees your documents. OCR'd rows
are lower-fidelity; the caller flags them for review.
"""

from __future__ import annotations

import csv
import io
import shutil
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except (AttributeError, ValueError):
    pass


def read_any(path: str, sheet=None):
    """Dispatch on extension -> (rows, note). `sheet` selects the worksheet for .xlsx/.xlsm."""
    ext = Path(path).suffix.lower()
    if ext in (".xlsx", ".xlsm"):
        return read_xlsx(path, sheet=sheet)
    if ext in (".csv", ".tsv", ".txt"):
        return read_csv(path)
    if ext == ".pdf":
        return read_pdf(path)
    if ext in (".docx",):
        return read_docx(path)
    if ext == ".msg":
        return read_msg(path)
    raise ValueError(f"Unsupported source type: {ext or '(no extension)'}")


class SheetSelectionRequired(ValueError):
    """Raised when an .xlsx has several non-empty sheets and none was specified — real
    workbooks carry cover sheets, multiple tabs and exports, so silently reading the 'active'
    sheet is a foot-gun. The caller should show `sheets` and re-call with `sheet=<name>`."""
    def __init__(self, path, sheets):
        self.path = path
        self.sheets = sheets
        listing = ", ".join(
            f"'{s['name']}' ({s['rows']}×{s['cols']}{', hidden' if s['hidden'] else ''})"
            for s in sheets)
        super().__init__(
            f"{Path(path).name} has {len(sheets)} non-empty sheets: {listing}. "
            f"Specify one, e.g. read_any(path, sheet='{sheets[0]['name']}').")


def _sheet_has_data(ws):
    """True as soon as any cell holds a non-blank value (cheap for non-empty sheets)."""
    for row in ws.iter_rows(values_only=True):
        if any(c is not None and str(c).strip() != "" for c in row):
            return True
    return False


def list_sheets(path: str):
    """Discover the worksheets in a workbook -> [{name, rows, cols, hidden, active, empty}],
    so the caller (or the user) can choose which to read."""
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    try:
        active_title = wb.active.title if wb.active is not None else None
        out = []
        for ws in wb.worksheets:
            out.append({"name": ws.title,
                        "rows": ws.max_row or 0, "cols": ws.max_column or 0,
                        "hidden": ws.sheet_state != "visible",
                        "active": ws.title == active_title,
                        "empty": not _sheet_has_data(ws)})
        return out
    finally:
        wb.close()


def read_xlsx(path: str, sheet=None):
    """Read one worksheet -> (rows, note). With `sheet=None`: auto-select the single non-empty
    visible sheet; if several are non-empty, raise SheetSelectionRequired (never guess via the
    'active' sheet). Pass `sheet=<name>` to read a specific tab."""
    from openpyxl import load_workbook
    sheets = list_sheets(path)
    by_name = {s["name"]: s for s in sheets}
    auto = False
    if sheet is not None:
        if sheet not in by_name:
            raise ValueError(f"Sheet {sheet!r} not found in {Path(path).name}. "
                             f"Available: {', '.join(by_name) or '(none)'}")
        target = sheet
    else:
        candidates = [s for s in sheets if not s["empty"] and not s["hidden"]]
        if len(candidates) == 1:
            target, auto = candidates[0]["name"], True
        elif not candidates:
            non_empty = [s for s in sheets if not s["empty"]]   # all empty/hidden — best effort
            target = (non_empty[0]["name"] if non_empty
                      else next((s["name"] for s in sheets if s["active"]),
                                sheets[0]["name"] if sheets else "Sheet1"))
        else:
            raise SheetSelectionRequired(path, candidates)
    wb = load_workbook(path, data_only=True, read_only=True)
    try:
        ws = wb[target]
        rows = [["" if c is None else c for c in row]
                for row in ws.iter_rows(values_only=True)]
    finally:
        wb.close()
    note = f"xlsx '{target}', {len(rows)} raw rows"
    if auto:
        note += " (auto-selected the only data sheet)"
    return rows, note


def read_csv(path: str):
    with open(path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
        except csv.Error:
            dialect = csv.excel
        rows = [list(r) for r in csv.reader(f, dialect)]
    return rows, f"delimited text, {len(rows)} raw rows"


def read_paste(text: str):
    """Parse a pasted table: markdown pipe-table, else tab/multi-space/CSV-ish lines."""
    lines = [ln for ln in text.splitlines() if ln.strip() != ""]
    if not lines:
        return [], "empty paste"
    if all("|" in ln for ln in lines[:3]):  # markdown table
        rows = []
        for ln in lines:
            if set(ln.strip()) <= set("|-: "):  # separator row
                continue
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            rows.append(cells)
        return rows, f"pasted markdown table, {len(rows)} rows"
    if "\t" in text:
        return [ln.split("\t") for ln in lines], f"pasted TSV, {len(lines)} rows"
    try:
        rows = [list(r) for r in csv.reader(io.StringIO(text))]
        if max(len(r) for r in rows) > 1:
            return rows, f"pasted CSV, {len(rows)} rows"
    except csv.Error:
        pass
    import re
    return [re.split(r"\s{2,}", ln.strip()) for ln in lines], f"pasted (split on 2+ spaces), {len(lines)} rows"


# --------------------------------------------------------------------------- #
# PDF table extraction — two engines, best result wins per page.
#
# PDF tables are messy by default, so we don't trust one extractor. pdfplumber (built on
# pdfminer.six, finely configurable: lines vs whitespace strategies) is PREFERRED for digital
# tables — it's usually better on borderless / whitespace-aligned tables. PyMuPDF
# (`find_tables`) is fast, strong on RULED tables, and is the ONLY engine that can drive OCR
# for scanned pages. So per page we run both (when both are installed), SCORE the result, and
# keep the better one; scanned pages always go through PyMuPDF + local Tesseract. pdfplumber is
# an OPTIONAL dependency — absent, behaviour falls back to PyMuPDF exactly as before.
# --------------------------------------------------------------------------- #
def _norm_table(t):
    return [["" if c is None else c for c in r] for r in t]


def _table_score(rows):
    """Score a candidate table: cells in its most-consistent column block. Penalises 1-column
    'tables' (really text) and one-row fragments. Higher = a more table-shaped extraction."""
    from collections import Counter
    rows = [r for r in rows if any(str(c).strip() for c in r)]
    if len(rows) < 2:
        return 0
    widths = Counter(len(r) for r in rows)
    mode_w, mode_n = widths.most_common(1)[0]
    if mode_w < 2:
        return 0
    return mode_n * mode_w


def _choose_tables(pp, mu):
    """Pick the better-scoring table set for a page -> (tables, engine). Ties go to pdfplumber
    (the preferred extractor for messy tables)."""
    ps, ms = sum(_table_score(t) for t in pp), sum(_table_score(t) for t in mu)
    if ps == 0 and ms == 0:
        return [], None
    if pp and ps >= ms:
        return pp, "pdfplumber"
    return mu, "pymupdf"


def _pp_page_tables(page):
    """pdfplumber tables for one page: try the ruled (lines) strategy then whitespace (text),
    keep the higher-scoring. -> list of tables (each a list of rows)."""
    best = None
    for settings in ({"vertical_strategy": "lines", "horizontal_strategy": "lines"},
                     {"vertical_strategy": "text", "horizontal_strategy": "text"}):
        try:
            tbls = page.extract_tables(settings)
        except Exception:  # noqa: BLE001 — strategy can throw on odd pages
            tbls = None
        if not tbls:
            continue
        norm = [_norm_table(t) for t in tbls]
        score = sum(_table_score(t) for t in norm)
        if best is None or score > best[0]:
            best = (score, norm)
    return best[1] if best else []


def _pdfplumber_tables_by_page(path):
    """{page_no: [tables...]} via pdfplumber, or None if pdfplumber isn't installed."""
    try:
        import pdfplumber
    except ImportError:
        return None
    out = {}
    try:
        with pdfplumber.open(path) as pdf:
            for pno, page in enumerate(pdf.pages):
                tbls = _pp_page_tables(page)
                if tbls:
                    out[pno] = tbls
    except Exception:  # noqa: BLE001 — degrade to PyMuPDF rather than crash
        return None
    return out


def _mupdf_page_tables(page):
    """PyMuPDF tables for one already-open page -> list of tables (each a list of rows)."""
    out = []
    try:
        for tbl in page.find_tables().tables:
            ext = tbl.extract()
            if ext:
                out.append(_norm_table(ext))
    except Exception:  # noqa: BLE001 — find_tables can be brittle on odd PDFs
        pass
    return out


def _mupdf_tables_by_page(path):
    """{page_no: [tables...]} via PyMuPDF, or {} if PyMuPDF isn't installed."""
    try:
        import fitz
    except ImportError:
        return {}
    out = {}
    for pno, page in enumerate(fitz.open(path)):
        tbls = _mupdf_page_tables(page)
        if tbls:
            out[pno] = tbls
    return out


def read_pdf(path: str):
    """Digital PDF tables via pdfplumber + PyMuPDF (best result per page); text-layout split for
    columnar pages with no detected table; OCR fallback (local Tesseract) for scanned pages."""
    try:
        import fitz  # PyMuPDF — backbone for text + OCR
    except ImportError:
        return [], "PyMuPDF (fitz) not installed — cannot read PDF"
    import re
    pp_by_page = _pdfplumber_tables_by_page(path)        # None if pdfplumber absent
    rows, ocr_pages, scanned, text_pages = [], 0, 0, 0
    pp_used, mu_used = 0, 0
    for pno, page in enumerate(fitz.open(path)):
        pp = (pp_by_page or {}).get(pno, [])
        mu = _mupdf_page_tables(page)
        chosen, eng = _choose_tables(pp, mu)
        if chosen:
            for t in chosen:
                rows.extend(t)
            if eng == "pdfplumber":
                pp_used += 1
            else:
                mu_used += 1
            continue
        text = page.get_text("text")
        if text.strip():                          # digital, columnar text but no detected table
            rows.extend([re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()])
            text_pages += 1
        else:                                     # no text layer -> likely scanned
            scanned += 1
            ocr_rows = _ocr_page(page)
            if ocr_rows:
                rows.extend(ocr_rows)
                ocr_pages += 1
    note = f"PDF, {len(rows)} rows"
    eng_bits = [b for b in (f"{pp_used} via pdfplumber" if pp_used else "",
                            f"{mu_used} via PyMuPDF" if mu_used else "") if b]
    if eng_bits:
        note += f"; tables: {', '.join(eng_bits)}"
    if pp_by_page is None and (text_pages or mu_used):
        note += " (pdfplumber not installed — install it for messy/borderless tables)"
    if text_pages:
        note += f"; {text_pages} page(s) via text-layout split (check column alignment)"
    if scanned:
        note += (f"; {ocr_pages}/{scanned} scanned page(s) OCR'd (LOWER FIDELITY — review)"
                 if ocr_pages else f"; {scanned} scanned page(s) need OCR — Tesseract not available")
    return rows, note


def _ocr_page(page):
    """OCR one PyMuPDF page via local Tesseract, returning whitespace-split rows. [] if no engine."""
    if not ocr_available():
        return []
    try:
        tp = page.get_textpage_ocr(flags=0, full=True)  # uses local Tesseract
        text = page.get_text("text", textpage=tp)
    except Exception:  # noqa: BLE001
        return []
    import re
    return [re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()]


def ocr_available() -> bool:
    """Local Tesseract present? (binary on PATH + PyMuPDF OCR support). Cloud OCR is never used."""
    if not shutil.which("tesseract"):
        return False
    try:
        import fitz  # noqa: F401
        return hasattr(fitz.Page, "get_textpage_ocr")
    except ImportError:
        return False


def read_docx(path: str):
    try:
        from docx import Document
    except ImportError:
        return [], "python-docx not installed — cannot read .docx tables"
    rows = []
    doc = Document(path)
    for tbl in doc.tables:
        for row in tbl.rows:
            rows.append([c.text for c in row.cells])
    return rows, f"docx, {len(rows)} table rows from {len(doc.tables)} table(s)"


def read_msg(path: str):
    try:
        import extract_msg
    except ImportError:
        return [], "extract_msg not installed — cannot read Outlook .msg"
    body = extract_msg.Message(path).body or ""
    return read_paste(body)[0], "Outlook .msg body parsed as a pasted table"


# --------------------------------------------------------------------------- #
# Document -> plain text + table enumeration (for data-extract: key-value
# field extraction and multi-table selection)
# --------------------------------------------------------------------------- #
def read_text(path: str):
    """Whole-document plain text -> (text, note). PDF uses the text layer with a local-OCR
    fallback for scanned pages; .docx joins paragraphs + table cells; .msg the body; .txt raw."""
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        try:
            import fitz
        except ImportError:
            return "", "PyMuPDF (fitz) not installed — cannot read PDF text"
        parts, ocr_pages, scanned = [], 0, 0
        for page in fitz.open(path):
            t = page.get_text("text")
            if t.strip():
                parts.append(t)
            elif ocr_available():
                try:
                    tp = page.get_textpage_ocr(flags=0, full=True)
                    parts.append(page.get_text("text", textpage=tp))
                    ocr_pages += 1
                except Exception:  # noqa: BLE001
                    scanned += 1
            else:
                scanned += 1
        note = "PDF text"
        if ocr_pages:
            note += f"; {ocr_pages} page(s) OCR'd (review)"
        if scanned:
            note += f"; {scanned} scanned page(s) unreadable (no Tesseract)"
        return "\n".join(parts), note
    if ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            return "", "python-docx not installed"
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                lines.append("\t".join(c.text for c in row.cells))
        return "\n".join(lines), "docx text"
    if ext == ".msg":
        try:
            import extract_msg
        except ImportError:
            return "", "extract_msg not installed"
        return (extract_msg.Message(path).body or ""), "msg body"
    if ext in (".txt", ".csv", ".tsv"):
        return Path(path).read_text(encoding="utf-8-sig", errors="replace"), "text file"
    raise ValueError(f"read_text: unsupported type {ext or '(none)'}")


def list_pdf_tables(path: str):
    """Enumerate tables per page -> [{page, index, engine, rows, cols, preview}] so the caller
    can pick which to extract. Per page the better-scoring engine (pdfplumber vs PyMuPDF) wins;
    `engine` records which produced it (pass it back to `extract_pdf_table` for an exact pull)."""
    pp = _pdfplumber_tables_by_page(path) or {}
    mu = _mupdf_tables_by_page(path)
    out = []
    for pno in sorted(set(pp) | set(mu)):
        chosen, eng = _choose_tables(pp.get(pno, []), mu.get(pno, []))
        for i, t in enumerate(chosen):
            if not t:
                continue
            out.append({"page": pno, "index": i, "engine": eng, "rows": len(t),
                        "cols": max(len(r) for r in t), "preview": t[:2]})
    return out


def extract_pdf_table(path: str, page: int, index: int = 0, engine=None):
    """Pull one specific table (by page + index from list_pdf_tables) -> rows. `engine`
    ('pdfplumber' | 'pymupdf') forces a specific extractor; default re-picks the best for that
    page (consistent with list_pdf_tables)."""
    pp = (_pdfplumber_tables_by_page(path) or {}).get(page, [])
    mu = _mupdf_tables_by_page(path).get(page, [])
    if engine == "pdfplumber":
        chosen = pp
    elif engine == "pymupdf":
        chosen = mu
    else:
        chosen, _ = _choose_tables(pp, mu)
    if index < 0 or index >= len(chosen):
        raise IndexError(f"extract_pdf_table: requested table index {index}, "
                         f"but page {page} has {len(chosen)} table(s)")
    return chosen[index]


if __name__ == "__main__":
    paste = "Investor | Commit | Close\nAcme Pension | £1,000,000 | 12/06/2026\nBeta FO | S$2.5m | 2026-06-13"
    rows, note = read_paste(paste)
    print("[paste]", note)
    for r in rows:
        print("  ", r)
    print("[ocr] local Tesseract available:", ocr_available())
