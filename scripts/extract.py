"""data-extract — get STRUCTURED data OUT of documents (PDF/scan/Word/email), then hand
to the shared cleanup engine. The counterpart to data-tidy: tidy cleans already-tabular
data; extract locates and pulls data from document-shaped sources —

  - multi-table / multi-page docs  -> list_tables() to pick, get_table() to pull
  - key-value / form / certificate -> extract_fields() (label -> value, one record)
  - scanned docs                   -> via shared ingest's local-OCR fallback

A SHARED module at the toolkit-root scripts/ (with dataclean.py + ingest.py); used by the
data-extract skill and by generated reuse runners (see dataclean.emit_runner). Output +
normalisation + the change report all come from dataclean.

    import sys, pathlib
    sys.path.insert(0, str(pathlib.Path("../../scripts").resolve()))   # shared engine
    import extract
    rec, flags = extract.extract_fields("certificate.pdf", FIELDS)     # form -> one record
    cands = extract.list_tables("report.pdf")                          # which tables exist
    rows  = extract.get_table("report.pdf", page=1, index=0)           # pull one

DATA HANDLING: this engine runs on your machine and makes no network calls (local OCR only —
never a cloud OCR). But the AI agent driving it sends whatever it reads into its context to your
AI provider — "never leaves the machine" is NOT claimed. See ../PRINCIPLES.md (§ Data handling).
"""

from __future__ import annotations

import json
import re
import sys
import pathlib
from decimal import Decimal

try:
    sys.stdout.reconfigure(encoding="utf-8")
except (AttributeError, ValueError):
    pass

# dataclean + ingest live in the same (shared) scripts/ dir as this file
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
import dataclean  # noqa: E402
import ingest  # noqa: E402


_FIELD_SPEC_CACHE = {}


def _load_field_spec(fields, source=None):
    """Resolve a field list supplied inline or as a JSON path.

    Schema validation records resolved specs in a context-local cache, so normal
    validate-then-run execution remains plan-relative without mutating the plan.
    Direct engine calls can use an absolute path, a cwd-relative path, or a path
    found relative to the source document or one of its parent directories.
    """
    if not isinstance(fields, (str, pathlib.Path)):
        return fields

    raw = str(fields)
    if raw in _FIELD_SPEC_CACHE:
        return _FIELD_SPEC_CACHE[raw]
    try:
        import agent_schemas
        cached = agent_schemas.cached_spec_value(raw)
    except (ImportError, AttributeError):
        cached = None
    if cached is not None:
        _FIELD_SPEC_CACHE[raw] = cached
        return cached

    path = pathlib.Path(fields).expanduser()
    candidates = [path] if path.is_absolute() else [pathlib.Path.cwd() / path]
    if source is not None and not path.is_absolute():
        source_path = pathlib.Path(source).expanduser()
        if not source_path.is_absolute():
            source_path = (pathlib.Path.cwd() / source_path).resolve()
        for parent in (source_path.parent, *source_path.parents):
            candidate = parent / path
            if candidate not in candidates:
                candidates.append(candidate)

    for candidate in candidates:
        if candidate.is_file():
            value = json.loads(candidate.read_text(encoding="utf-8-sig"))
            if not isinstance(value, list):
                raise ValueError(f"field spec must be a JSON list: {candidate}")
            _FIELD_SPEC_CACHE[raw] = value
            _FIELD_SPEC_CACHE[str(candidate.resolve())] = value
            return value
    attempted = ", ".join(str(candidate) for candidate in candidates)
    raise FileNotFoundError(f"field spec not found: {raw} (tried {attempted})")


# --------------------------------------------------------------------------- #
# Key-value / form extraction  ->  one record per document
# --------------------------------------------------------------------------- #
def extract_fields(source, fields, text=None):
    """Pull labelled fields from a document or text into one record.

    fields = [{"name": "Investor", "labels": ["investor", "name of investor"], "type": "text"},
              {"name": "Commitment", "labels": ["commitment", "amount"], "type": "currency",
               "currency": "GBP"},
              {"name": "Close date", "labels": ["close date", "closing"], "type": "date"}]
    Returns (record: {name: value}, flags: [{field, issue}]). Unfound/odd values are flagged,
    never invented. `source` may be a path (read via ingest.read_text) or pass `text=`.
    """
    fields = _load_field_spec(fields, source=source)
    if text is None:
        text, _ = ingest.read_text(str(source))
    lines = [ln.rstrip() for ln in text.splitlines()]
    all_labels = [lab for f in fields for lab in f.get("labels", [f["name"]])]
    record, flags = {}, []
    for f in fields:
        # a currency field may emit its detected code into its own column (see field_columns)
        ct = f.get("code_target") if f.get("type") == "currency" else None
        raw = _find_value(lines, f.get("labels", [f["name"]]), all_labels=all_labels)
        if raw is None:
            record[f["name"]] = ""
            if ct:
                record[ct] = ""
            flags.append({"field": f["name"], "issue": "label not found"})
            continue
        spec = {"type": f.get("type", "text"), "currency": f.get("currency"),
                "dayfirst": f.get("dayfirst", True)}
        val, note, kept_raw = dataclean.convert_value(raw, spec)
        record[f["name"]] = val
        if ct:                                       # keep the currency code (amount + code)
            record[ct] = dataclean.currency_code(raw, spec)
        if note:
            flags.append({"field": f["name"], "issue": note,
                          "value": raw, "kept_raw": kept_raw})
    return record, flags


def field_columns(fields):
    """Output column order for a field list — inserts a currency field's `code_target` column
    right after it, so `fields_to_table` keeps amount + code (mirrors the tidy recipe's
    code_target). Use: `extract.fields_to_table(records, extract.field_columns(FIELDS))`."""
    fields = _load_field_spec(fields)
    cols = []
    for f in fields:
        cols.append(f["name"])
        if f.get("type") == "currency" and f.get("code_target"):
            cols.append(f["code_target"])
    return cols


def _clean_value(s):
    """Tidy a captured value: drop box pipes and trailing dotted leaders, collapse spaces."""
    s = re.sub(r"\.{2,}\s*$", "", s.strip().strip("|").strip()).strip()
    return re.sub(r"\s{2,}", " ", s) or None


def _looks_like_label(line, others):
    """A line that is itself a field label — so a next-line search doesn't swallow the FOLLOWING
    field's label as this field's value. True if it matches a known other label or ends in ':'."""
    t = line.strip().lower().rstrip(":").strip()
    return t in others or line.strip().endswith(":")


def _find_value(lines, labels, *, all_labels=None):
    """First value for any of `labels`. Handles, in order:
      - same line:  'Label: value' · 'Label<tab>value' · 'Label  value' (2+ spaces) ·
                    'Label ....... value' (dotted leader)
      - next line:  the label sits alone on its line (optionally a trailing ':'), the value is
                    on the next non-empty line — common on confirmations / certificates / boxed
                    forms. The search stops if it hits the NEXT field's label (so it never
                    grabs a neighbouring label as this field's value).
    `all_labels` (every field's labels) powers the next-line guard. Returns the value or None."""
    norm_labels = [lab.strip() for lab in labels if lab and lab.strip()]
    others = ({lab.strip().lower() for lab in (all_labels or [])}
              - {lab.lower() for lab in norm_labels})
    n = len(lines)
    for label in norm_labels:
        lab = re.escape(label)
        for i, ln in enumerate(lines):
            m = (re.match(rf"\s*{lab}\s*[:\t]\s*(.+)$", ln, re.IGNORECASE)
                 or re.match(rf"\s*{lab}\s*\.{{2,}}\s*(.+)$", ln, re.IGNORECASE)
                 or re.match(rf"\s*{lab}\s{{2,}}(.+)$", ln, re.IGNORECASE))
            if m:
                cv = _clean_value(m.group(1))
                if cv:
                    return cv
            # label alone on its line (optionally trailing ':') -> value on a following line
            if re.fullmatch(rf"\s*{lab}\s*:?\s*", ln, re.IGNORECASE):
                for j in range(i + 1, min(i + 4, n)):
                    nxt = lines[j].strip()
                    if not nxt:
                        continue
                    if _looks_like_label(nxt, others):
                        break                      # ran into the next field — value is absent
                    cv = _clean_value(nxt)
                    if cv:
                        return cv
                    break
    return None


def fields_to_table(records, field_names):
    """Many per-document records -> (header, rows) for one combined .xlsx."""
    return list(field_names), [[r.get(n, "") for n in field_names] for r in records]


# --------------------------------------------------------------------------- #
# Table selection  ->  pick one of several tables across pages
# --------------------------------------------------------------------------- #
def list_tables(path: str):
    """Enumerate candidate tables in a document -> [{page,index,rows,cols,preview}].
    PDFs via PyMuPDF; .docx via python-docx (page=-1, index=table order)."""
    ext = pathlib.Path(path).suffix.lower()
    if ext == ".pdf":
        return ingest.list_pdf_tables(path)
    if ext == ".docx":
        from docx import Document
        out = []
        for i, tbl in enumerate(Document(path).tables):
            ext_rows = [[c.text for c in row.cells] for row in tbl.rows]
            out.append({"page": -1, "index": i, "rows": len(ext_rows),
                        "cols": max((len(r) for r in ext_rows), default=0),
                        "preview": ext_rows[:2]})
        return out
    return []


def get_table(path: str, page: int = 0, index: int = 0, engine=None):
    """Pull one specific table (from list_tables coordinates) -> rows."""
    ext = pathlib.Path(path).suffix.lower()
    if ext == ".pdf":
        return ingest.extract_pdf_table(path, page, index, engine=engine)
    if ext == ".docx":
        from docx import Document
        tbl = Document(path).tables[index]
        return [[c.text for c in row.cells] for row in tbl.rows]
    raise ValueError(f"get_table: unsupported type {ext or '(none)'}")


def render_fields_report(records, flags_list):
    """Brief report for a batch of extracted records + their flags."""
    out = [f"# Data-extract report — {len(records)} document(s)"]
    total = sum(len(f) for f in flags_list)
    out.append(f"- Flags across all documents: **{total}** (unfound / verify)")
    for i, flags in enumerate(flags_list):
        if flags:
            out.append(f"\n## ⚑ Document {i + 1}")
            out.append("| Field | Issue | Value |\n|---|---|---|")
            for fl in flags:
                out.append(f"| {_md_escape(fl['field'])} | {_md_escape(fl['issue'])} | "
                           f"{_md_escape(fl.get('value', ''))} |")
    return "\n".join(out)


def _md_escape(v):
    return str("" if v is None else v).replace("|", "\\|").replace("\r\n", "<br>").replace("\n", "<br>")


if __name__ == "__main__":
    FIELDS = [
        {"name": "Investor", "labels": ["investor", "name of investor"], "type": "text"},
        # code_target keeps the detected currency code beside the amount (mixed-currency batch)
        {"name": "Commitment", "labels": ["commitment", "amount"], "type": "currency",
         "code_target": "Commitment ccy"},
        {"name": "Close date", "labels": ["close date", "closing"], "type": "date"},
        {"name": "Settlement bank", "labels": ["settlement bank", "bank"], "type": "text"},
        {"name": "Fee", "labels": ["fee"], "type": "currency", "currency": "GBP"},
        {"name": "Reference", "labels": ["reference", "ref"], "type": "text"},
    ]
    sample = ("ABC Capital — Subscription confirmation\n"
              "Investor: Acme Pension Fund\n"
              "Commitment : S$ 2,000,000\n"             # a different currency to the fee
              "Close date\t12/06/2026\n"
              "Settlement bank\n"                       # label alone -> value on the next line
              "HSBC London\n"
              "Fee .......... GBP 2,500\n"              # dotted leader
              "(no reference quoted)\n")
    rec, flags = extract_fields(None, FIELDS, text=sample)
    print("[extract_fields] record:", rec)
    print("[extract_fields] flags:", flags)
    assert rec["Settlement bank"] == "HSBC London", rec       # next-line layout
    assert rec["Fee"] == Decimal("2500"), rec                 # dotted leader + currency
    assert rec["Commitment"] == Decimal("2000000"), rec
    assert rec["Commitment ccy"] == "SGD", rec                # code_target keeps the code
    cols = field_columns(FIELDS)
    assert cols[1:3] == ["Commitment", "Commitment ccy"], cols  # code column follows its amount
    hdr, rows = fields_to_table([rec], cols)
    assert hdr == cols and rows[0][2] == "SGD", (hdr, rows)
    print()
    print(render_fields_report([rec], [flags]))
